import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
from datetime import datetime, date, time
from io import BytesIO


def set_cell_border(cell, **kwargs):
    """Set cell border individually"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '4')
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), kwargs[edge])
            tcPr.append(element)


def get_year_from_semester(semester):
    try:
        semester = int(semester)
        if semester in [1, 2]:
            return 1
        elif semester in [3, 4]:
            return 2
        elif semester in [5, 6]:
            return 3
        elif semester in [7, 8]:
            return 4
    except:
        return None


def sort_and_organize_data(data):
    """Sort data by semester first, then by branch"""
    # Convert semester to int for proper numerical sorting
    data['Semester'] = data['Semester'].astype(int)
    
    # Sort by Semester first, then by Branch
    sorted_data = data.sort_values(['Semester', 'Branch'], ascending=[True, True])
    
    return sorted_data


def create_word_document(data):
    """Create Word document and return it as bytes"""
    doc = Document()
    
    # Sort and organize the data first
    sorted_data = sort_and_organize_data(data)
    
    # Group the sorted data
    grouped_data = sorted_data.groupby(['Semester', 'Branch'])
    
    # Process groups in the order they appear (which is now sorted)
    for (semester, branch), group in grouped_data:
        course_branch = str(branch)
        semester = int(semester)
        year = get_year_from_semester(semester)

        table = doc.add_table(rows=3, cols=7)
        table.autofit = False
        table.style = 'Table Grid'

        row0 = table.rows[0].cells
        row0[0].merge(row0[4]).text = f'Course: B.Tech {course_branch}'
        row0[5].merge(row0[6]).text = f'Year/Semester: {year}/{semester}'

        row1 = table.rows[1].cells
        headers = ['S.No', 'Name', 'Registration No.', 'Section', 'Date', 'Hours', '']
        for i, text in enumerate(headers):
            run = row1[i].paragraphs[0].add_run(text)
            run.bold = True

        row1[5].merge(row1[6])

        row2 = table.rows[2].cells
        subheaders = ['', '', '', '', '', 'From', 'To']
        for i, text in enumerate(subheaders):
            if text:
                run = row2[i].paragraphs[0].add_run(text)
                run.bold = True

        for i in range(5):
            table.cell(1, i).merge(table.cell(2, i))

        # First, apply borders to ALL rows (including first row)
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(cell, top='000000', left='000000', bottom='000000', right='000000')
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)

        # Then, specifically override the first row borders - set top, left, right to white and bottom to black
        for i in range(7):
            set_cell_border(row0[i], top='FFFFFF', left='FFFFFF', bottom='000000', right='FFFFFF')
        
        # Here is the section for the headings in the OD list
        for idx, (_, row) in enumerate(group.iterrows(), start=1):
            tr = table.add_row().cells
            tr[0].text = str(idx)
            tr[1].text = str(row['Name'])
            tr[2].text = str(row['Registration Number'])
            tr[3].text = str(row['Section'])
            tr[4].text = str(row['Date'])  # This now comes from user input
            tr[5].text = str(row['From']) if pd.notna(row['From']) else ''
            tr[6].text = str(row['To']) if pd.notna(row['To']) else ''

            if idx == 1:
                for cell in tr:
                    set_cell_border(cell, top='FFFFFF', left='000000', bottom='000000', right='000000')
            else:
                for cell in tr:
                    set_cell_border(cell, top='000000', left='000000', bottom='000000', right='000000')

            for cell in tr:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)

        doc.add_paragraph()
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run("Event Coordinator").bold = True
        p.add_run("\t\t\t")
        p.add_run("Head Student Welfare").bold = True
        p.add_run("\t\t\t")
        p.add_run("HOD").bold = True

    # Save to BytesIO instead of file
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer.getvalue()


def process_uploaded_data(uploaded_file, event_date, from_time, to_time):
    """Process uploaded CSV/Excel file and add user-provided date and times"""
    try:
        # Read uploaded file
        if uploaded_file.name.endswith('.csv'):
            data = pd.read_csv(uploaded_file)
        elif uploaded_file.name.endswith(('.xlsx', '.xls')):
            data = pd.read_excel(uploaded_file)
        else:
            raise ValueError("Unsupported file format")
            
        # Updated required columns (removed Date, From, To since user provides them)
        required_cols = ['Name', 'Registration Number', 'Section', 'Branch', 'Semester']
        missing_cols = [col for col in required_cols if col not in data.columns]
        
        if missing_cols:
            raise ValueError(f"Missing columns: {', '.join(missing_cols)}")
        
        # Add user-provided date and times to all rows
        data['Date'] = event_date.strftime("%d-%m-%Y")
        data['From'] = from_time.strftime("%H:%M")
        data['To'] = to_time.strftime("%H:%M")
            
        return data
        
    except Exception as e:
        st.error(f"Error processing file: {e}")
        return pd.DataFrame()


def main():
    # Page configuration
    st.set_page_config(
        page_title="OD List Generator",
        page_icon="📋",
        layout="wide"
    )
    
    # Title and description
    st.title("📋 OD List Generator")
    st.markdown("Generate Official Duty (OD) lists in Word format from your participant data")
    
    # Sidebar for instructions
    with st.sidebar:
        st.header("📖 Instructions")
        st.markdown("""
        **Required Columns in your file:**
        - Name
        - Registration Number
        - Section
        - Branch
        - Semester
        
        **Note:** Date, From, and To times will be entered on the main page and applied to all participants.
        
        **Supported formats:**
        - CSV (.csv)
        - Excel (.xlsx, .xls)
        """)
        
        # Sample data download (updated)
        st.header("📥 Sample Data")
        sample_data = {
            'Name': ['John Doe', 'Jane Smith', 'Mike Johnson'],
            'Registration Number': ['REG001', 'REG002', 'REG003'],
            'Section': ['A', 'B', 'A'],
            'Branch': ['CSE', 'ECE', 'CSE'],
            'Semester': [1, 2, 1]
        }
        sample_df = pd.DataFrame(sample_data)
        csv_sample = sample_df.to_csv(index=False)
        
        st.download_button(
            label="Download Sample CSV",
            data=csv_sample,
            file_name="sample_od_data.csv",
            mime="text/csv"
        )
    
    # Main content area
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.header("📤 Upload Your Data")
        
        # File uploader
        uploaded_file = st.file_uploader(
            "Choose a file",
            type=['csv', 'xlsx', 'xls'],
            help="Upload CSV or Excel file with participant data"
        )
        
        # Event details section
        st.header("📅 Event Details")
        
        col_date, col_from, col_to = st.columns([1, 1, 1])
        
        with col_date:
            event_date = st.date_input(
                "Event Date",
                value=date.today(),
                help="Select the date for the OD"
            )
        
        with col_from:
            from_time = st.time_input(
                "From Time",
                value=time(9, 0),  # Default 9:00 AM
                help="Start time for the OD"
            )
        
        with col_to:
            to_time = st.time_input(
                "To Time", 
                value=time(17, 0),  # Default 5:00 PM
                help="End time for the OD"
            )
        
        # Show selected details
        st.info(f"📋 **Event Details:** {event_date.strftime('%d-%m-%Y')} from {from_time.strftime('%H:%M')} to {to_time.strftime('%H:%M')}")
        
        if uploaded_file is not None:
            # Show file details
            st.success(f"✅ File uploaded: {uploaded_file.name}")
            
            # Validate time input
            if from_time >= to_time:
                st.error("❌ 'From Time' must be earlier than 'To Time'")
                return
            
            # Process the uploaded file
            with st.spinner("Processing file..."):
                df = process_uploaded_data(uploaded_file, event_date, from_time, to_time)
            
            if not df.empty:
                st.success(f"🎉 Data loaded successfully! **{len(df)} records** found.")
                
                # Show data preview
                st.subheader("📊 Data Preview")
                st.dataframe(df.head(10), use_container_width=True)
                
                # Show data statistics
                with col2:
                    st.header("📈 Data Statistics")
                    st.metric("Total Records", len(df))
                    st.metric("Unique Branches", df['Branch'].nunique())
                    st.metric("Unique Semesters", df['Semester'].nunique())
                    
                    # Show event details in stats
                    st.subheader("Event Details")
                    st.write(f"**Date:** {event_date.strftime('%d-%m-%Y')}")
                    st.write(f"**From:** {from_time.strftime('%H:%M')}")
                    st.write(f"**To:** {to_time.strftime('%H:%M')}")
                    
                    # Show branch-wise count
                    st.subheader("Branch Distribution")
                    branch_counts = df['Branch'].value_counts()
                    for branch, count in branch_counts.items():
                        st.write(f"**{branch}:** {count}")
                    
                    # Show semester-wise count
                    st.subheader("Semester Distribution")
                    semester_counts = df['Semester'].value_counts().sort_index()
                    for semester, count in semester_counts.items():
                        st.write(f"**Semester {semester}:** {count}")
                
                # Generate document section
                st.markdown("---")
                st.header("📄 Generate Document")
                
                col_gen1, col_gen2, col_gen3 = st.columns([1, 2, 1])
                
                with col_gen2:
                    if st.button("🚀 Generate OD Document", type="primary", use_container_width=True):
                        with st.spinner("Generating Word document... This may take a few moments."):
                            try:
                                # Generate timestamp for filename
                                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                                output_filename = f'OD_List_{event_date.strftime("%Y%m%d")}_{timestamp}.docx'
                                
                                # Create document
                                doc_bytes = create_word_document(df)
                                
                                # Success message
                                st.success("🎉 Document generated successfully!")
                                
                                # Download button
                                st.download_button(
                                    label="📥 Download OD Document",
                                    data=doc_bytes,
                                    file_name=output_filename,
                                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                                    use_container_width=True
                                )
                                
                            except Exception as e:
                                st.error(f"❌ Error generating document: {str(e)}")
                                st.exception(e)
            
            else:
                st.error("❌ Failed to process the uploaded file. Please check the format and required columns.")
        
    # Show upload instructions when no file is uploaded
    if uploaded_file is None:
        st.info("👆 Please upload a CSV or Excel file and set the event details to get started")
        
        # Show format example
        st.subheader("📋 Expected File Format")
        st.markdown("Your file should contain the following columns:")
        
        example_data = {
            'Name': ['Student Name'],
            'Registration Number': ['REG001'],
            'Section': ['A'],
            'Branch': ['CSE'],
            'Semester': [1]
        }
        example_df = pd.DataFrame(example_data)
        st.dataframe(example_df, use_container_width=True)
        
        st.markdown("**Note:** Date, From, and To times will be entered above and applied to all participants.")


if __name__ == "__main__":
    main()
